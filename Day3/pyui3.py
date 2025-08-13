import tkinter as tk
from os.path import basename
from tkinter import ttk,messagebox
import os.path
import shutil
import csv
from docx import Document
from tkinter import filedialog

#main window
root = tk.Tk()
root.geometry("440x400")
root.title("Palindrome Check")
root.resizable(True, True)

def verify_user_entry(event):
    reverse_label.config(text="")
    if not event.char.isalnum():
        given_string_entry.delete(0, tk.END)
        given_string_entry.insert(0, given_string_entry.get()[:-1])

def show_reverse():
    given_string = given_string_entry.get()
    if len(given_string.strip()) > 0:
        result_str = given_string[::-1]
        reverse_label.config(text=result_str)
        record_data(privilege.get())
    else:
        messagebox.showerror("Error", "Please enter a string")

def clear_all():
    given_string_entry.delete(0, tk.END)
    reverse_label.config(text="")

def record_data(user_type):
    #Record data when the selected preferance is administrator
    if user_type==1:
        given_string=given_string_entry.get()
        reversed_string=given_string[::-1]
        if given_string==reversed_string:
            with open(txt_file,"a") as fp:
                fp.write(given_string + "\n")
                doc = Document(doc_file)
                doc.add_paragraph(given_string).bold=True
                doc.save(doc_file)
        else:
            with open(csv_file,"a") as fp:
                writer = csv.writer(fp)
                writer.writerow([given_string, reversed_string])

def download_doc():
    result_doc=filedialog.asksaveasfile(
        mode="w",
        defaultextension=".docx",
        filetypes=(("DOCX", "*.docx"),("All files", "*"))
    )
    if result_doc is None:
        return
    name=result_doc.name
    basename=os.path.basename(name)
    path=os.path.dirname(name)
    target=os.path.join(path,basename)
    shutil.copy(doc_file,target)


txt_file='PalindromeText.txt'
csv_file='ReversedTable.csv'
doc_file='ReversedDocument.docx'

#create the palimdrome doc file
if not os.path.exists(doc_file):
    document=Document()
    document.add_heading("PALINDROME_CHECK",0)
    document.save(doc_file)

#create the palindromeText.txt file
if not os.path.exists(txt_file):
    with open(txt_file, 'w', newline='') as fp:
        fp.write('Palindrome List'+'\n'+'================'+'\n')

#create the reversedtable.csv file
if not os.path.exists(csv_file):
    with open(csv_file, 'w', newline='') as fp:
        writer = csv.writer(fp)
        writer.writerow(["Given", "Reversed"])

privilege = tk.IntVar()

pal_label=ttk.Label(root, text="Palindrome Check")
pal_label.grid(row=0, column=0,sticky=tk.W, columnspan=2, padx=20, pady=20)

admin_rdo_btn=ttk.Radiobutton(root, text="Administrator", variable=privilege, value=1,command=clear_all)
admin_rdo_btn.grid(row=1, column=0, sticky=tk.W, padx=20, pady=20)

user_rdo_btn=ttk.Radiobutton(root, text="User", variable=privilege, value=0,command=clear_all)
user_rdo_btn.grid(row=1, column=1, sticky=tk.W, padx=20, pady=20)

given_string_label = tk.Label(root, text="GIVEN STRING")
given_string_label.grid(column=0, row=2, sticky=tk.W,padx=5, pady=5)

given_string_entry = ttk.Entry(root)
given_string_entry.grid(column=1, row=2, sticky=tk.W,padx=5, pady=5)
given_string_entry.bind('<KeyRelease>', verify_user_entry)

reverse_btn = tk.Button(root, text="REVERSE", command=show_reverse)
reverse_btn.grid(column=2, row=2, sticky=tk.W,padx=5, pady=5)

result_label = tk.Label(root, text="REVERESED STRING:")
result_label.grid(column=0, row=3, sticky=tk.W,padx=5, pady=5)

reverse_label = tk.Label(root, text="",font='helvetica 25 bold')
reverse_label.grid(column=1, row=3, sticky=tk.W,padx=5, pady=5,columnspan=3)

download_label = tk.Label(root, text="DOWNLOAD")
download_label.grid(column=0, row=4, sticky=tk.W,padx=5, pady=5)

download_txt_btn = ttk.Button(root, text="Download CSV")
download_txt_btn.grid(column=1, row=4, sticky=tk.W, padx=10, pady=10)

download_doc_btn = ttk.Button(root, text="Download DOC",command=download_doc )
download_doc_btn.grid(column=2, row=4, sticky=tk.W, padx=10, pady=10)


root.mainloop()